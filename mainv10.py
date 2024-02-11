# Standard library imports
import os
import sys
import uuid
import re
import json
import logging
import hashlib
import glob
import shutil
from datetime import datetime
from operator import itemgetter
from html import escape
from urllib.parse import quote

# Data handling
import pandas as pd
import numpy as np

# PDF processing
import pdfplumber

# Document processing
from docx import Document

# image processing
from PIL import Image
import pytesseract

# Machine Learning and NLP
import torch
from transformers import (
    AutoModelForCausalLM, 
    AutoTokenizer, 
    pipeline, 
    BitsAndBytesConfig,
    AutoConfig
)
from langchain import (
    HuggingFacePipeline,
    PromptTemplate, 
    LLMChain
)

# OpenAI LLM
from openllm import LLM

# Streamlit for web apps
import streamlit as st
import streamlit.components.v1 as components

# ChromaDB utilities
import chromadb
from chromadb.utils import embedding_functions

# ahocorasick for the patten matches

import ahocorasick

# Langchain utilities
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Set the environment variable
os.environ['ALLOW_RESET'] = 'TRUE'

PYTESTACT_PATH = r'C:\Users\hp\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
BASE_FOLDER = 'C:\\contracts'
topics_excel_file_path = 'Categorized_Contract_Topics.xlsx'
full_excel_path = os.path.join(BASE_FOLDER, topics_excel_file_path)
st.write(f"Full Excel Path: {full_excel_path}")

def get_topics_from_excel(topics_excel_file_path, sheet_name='MASTER_TOPIC', column_name='CONTRACT_TOPIC'):
    """
    Reads topics from an Excel file.

    Args:
        excel_file_path (str): Path to the Excel file.
        sheet_name (str): Name of the sheet in the Excel file. Default is 'Sheet1'.
        column_name (str): Name of the column containing topics. Default is 'TOPIC'.

    Returns:
        set: A set of unique topics.
    """
    try:
        # Read the specified column from the Excel sheet
        df = pd.read_excel(topics_excel_file_path, sheet_name=sheet_name)
        st.write(df)
        topic_column = df[column_name]

        # Get unique topics
        unique_topics = set(topic_column.unique())

        return unique_topics
    except Exception as e:
        st.write(f"Error reading topics from Excel: {str(e)}")
        return set()  # Return an empty set in case of an error


# Find topics in the SEAL database
 
topics_to_find = get_topics_from_excel(full_excel_path)

if topics_to_find:
    st.write(topics_to_find)
else:
    st.write('Unable to get topics_to_find')

# Constants

NUMBER_SEARCH_QUERY = 1
MAX_DISPLAY_REQ = 1
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
VECTORDB_NAME_SPACE = 'ContractsDB'
EXTRACT_TABLE_STRATEGY = {"vertical_strategy": "lines_strict", "horizontal_strategy": "lines"}

#-----------------------------------------------------------------------------------
#-----------------------------------------------------------------------------------
# Initialize text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
    is_separator_regex=False
)

#-----------------------------------------------------------------------------------
# Function to create a download button
def create_download_button(file_path):
    with open(file_path, "rb") as file:
        btn = st.download_button(
                label="Download File",
                data=file,
                file_name=os.path.basename(file_path),
                mime="text/plain"
            )

LOG_FOLDER = os.path.join(BASE_FOLDER, 'LOG_FOLDER')
PDF_PREPROCESSED_LOG_FILE_PATH =  os.path.join(LOG_FOLDER, 'GenAIlog.log')
MODEL_FOLDER = os.path.join(BASE_FOLDER, 'modelslist')
DEFAULT_MODEL = "all-MiniLM-L6-v2"
DEFAULT_MODEL_PATH = os.path.join(MODEL_FOLDER, DEFAULT_MODEL)
UPLOAD_FILE_STATS = os.path.join(LOG_FOLDER, 'Chroma_Upload_PDF_Stats.xlsx')

# Create the LOG_FOLDER if it doesn't exist
if not os.path.exists(LOG_FOLDER):
    os.makedirs(LOG_FOLDER)
#-----------------------------------------------------------------------------------

PDF_FILES_TEMP_FOLDER = os.path.join(BASE_FOLDER, 'PDF_FILES_TEMP_FOLDER')

SMART_CHUNK_RESULT_OUTPUT_FOLDER = os.path.join(BASE_FOLDER, 'SMART_CHUNK_RESULT_OUTPUT_FOLDER')

PDF_OUTPUT_RESULT_FOLDER = os.path.join(BASE_FOLDER, 'PDF_OUTPUT_RESULT_FOLDER')

CHROMADB_PERSIST_FOLDER = os.path.join(BASE_FOLDER, 'CHROMADB_PERSIST_FOLDER')

ARCHIVE_FOLDER = os.path.join(BASE_FOLDER, 'ARCHIVE_FOLDER')

VECTORDB_DIR_SPACE_FOLDER = os.path.join(CHROMADB_PERSIST_FOLDER, VECTORDB_NAME_SPACE)

pdfdirectories = [BASE_FOLDER, PDF_FILES_TEMP_FOLDER, PDF_OUTPUT_RESULT_FOLDER, 
SMART_CHUNK_RESULT_OUTPUT_FOLDER, LOG_FOLDER, ARCHIVE_FOLDER,
CHROMADB_PERSIST_FOLDER, CHROMADB_PERSIST_FOLDER]
#-----------------------------------------------------------------------------------
# Ensure directories exist
for directory in pdfdirectories:
    if not os.path.exists(directory):
        os.makedirs(directory)
#-----------------------------------------------------------------------------------
STRATEGY_LIST = ['text', 'lines', 'lines_strict']
#-----------------------------------------------------------------------------------
# Generate combinations of strategies
EXTRACTION_STRATEGIES = {
    f"{v_strategy}_{h_strategy}": {"vertical_strategy": v_strategy, "horizontal_strategy": h_strategy}
    for v_strategy in STRATEGY_LIST for h_strategy in STRATEGY_LIST
}

# ------------------------------------------------------------------------
def get_model_directories():
    """Function to retrieve model directories."""
    model_dirs = [dir_name for dir_name in os.listdir(MODEL_FOLDER) if os.path.isdir(os.path.join(MODEL_FOLDER, dir_name))]
    return model_dirs

model_directories = get_model_directories()

# Check if default model exists in model directories; if not, prepend it
if DEFAULT_MODEL not in model_directories:
    model_directories.insert(0, DEFAULT_MODEL)  # Insert default model at the beginning

# Create model choices for Streamlit selectbox (use directory names directly)
model_choices = {dir_name: dir_name for dir_name in model_directories}

# Global debug flag
DEBUG_MODE = True
#-----------------------------------------------------------------------------------
def setup_logging(log_file):
    """
    Set up logging configurations.

    Args:
    - log_file (str): Name of the log file.

    Returns:
    - str: Full path to the log file.
    """
    SMART_LOG_PATH = os.path.join(LOG_FOLDER, log_file)

    logging.basicConfig(filename=SMART_LOG_PATH,
                        level=logging.INFO,
                        format='%(asctime)s - %(levelname)s - %(message)s')
    # Configure logging
    logging.basicConfig(filename=SMART_LOG_PATH,level=logging.DEBUG if DEBUG_MODE else logging.INFO,
                    format='%(asctime)s %(levelname)s: %(message)s',
                    datefmt='%Y-%m-%d %H:%M:%S')

    return SMART_LOG_PATH

# Setup logging and get the log file path
log_file_path = setup_logging(PDF_PREPROCESSED_LOG_FILE_PATH)
#-----------------------------------------------------------------------------------

def log_message(message, level=logging.INFO):
    """
    Log a given message and print if the log level isn't INFO.

    Args:
    - message (str): The message to log.
    - level (logging.LEVEL): The logging level.
    """
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    formatted_message = f"[{timestamp}] {message}"
    logging.log(level, formatted_message)
    if level != logging.INFO:
        print(formatted_message)


SEQUENCE_NUMBER = 1

def get_next_sequence():
    """
    Retrieve the next sequence number and increment the global sequence counter.
    
    Returns:
        int: The next sequence number.
    """
    global SEQUENCE_NUMBER
    val = SEQUENCE_NUMBER
    SEQUENCE_NUMBER += 1
    return val

def reset_sequence_number():
    """
    Reset the global sequence number to 1.
    """
    global SEQUENCE_NUMBER
    SEQUENCE_NUMBER = 1
#-----------------------------------------------------------------------------------
directories_to_archive = [
    PDF_OUTPUT_RESULT_FOLDER,
    SMART_CHUNK_RESULT_OUTPUT_FOLDER
]
#------------------------------------------------------------------
def archive_processed_files(GET_FILE_NAME):
    """
    Archive processed files to a specific directory.

    Args:
    - GET_FILE_NAME (str): The name of the file to be archived.

    This function moves processed files to an archive directory for each specified file.
    """
    for directory in directories_to_archive:
        arch_dir = os.path.join(ARCHIVE_FOLDER, GET_FILE_NAME, os.path.basename(directory))
        if not os.path.exists(arch_dir):
            os.makedirs(arch_dir)

        files_to_move = glob.glob(os.path.join(directory, '*'))
        for file_path in files_to_move:
            try:
                archive_file_path = os.path.join(arch_dir, os.path.basename(file_path))
                if not os.path.exists(archive_file_path):
                    shutil.move(file_path, arch_dir)
                else:
                    os.remove(file_path)
            except Exception as e:
                log_message(f"Error archiving {file_path} to {arch_dir}: {e}")

#-----------------------------------------------------------------------------------
ENGINE_LISTS = {
    'stage_2_extract_data_from_pdf': "Stage 1 PDF/Doc Extraction",
    'stage_3_chunking_process': "Stage 2 Data Chunking"
}
#-----------------------------------------------------------------------------------
# Utility Functions

def generate_key(base_string):
    """
    Generate a MD5 key from a string.

    Args:
    - base_string (str): The string to generate the key from.

    Returns:
    - str: MD5 hashed key.
    """
    return hashlib.md5(base_string.encode()).hexdigest()
#-----------------------------------------------------------------------------------
def delete_database(db_path):
    """
    Delete a database at a given path.

    Args:
    - db_path (str): The path to the database.

    This function deletes the database if it exists and logs success or error message.
    """
    if os.path.exists(db_path):
        shutil.rmtree(db_path)
        st.success("Database deleted successfully.")
    else:
        st.error("Database path does not exist.")
# -----------------------------------------------------------------
DELIMITER_TBL = '|'  # Define the delimiter if not already defined
# -----------------------------------------------------------------

def extract_data_from_img(img_file_path, output_path, file_name):
    """
    Extract data from an image file using Optical Character Recognition (OCR).

    Args:
        img_file_path (str): Path to the image file.
        output_path (str): Path to the output file.
        file_name (str): Name of the file being processed.

    Returns:
        list: Extracted data combined from the image file.
    """
    pytesseract.pytesseract.tesseract_cmd = PYTESTACT_PATH

    try:
        # Convert the uploaded file to an image object
        image = Image.open(img_file_path)

        # Using pytesseract to perform OCR on the image
        extracted_text = pytesseract.image_to_string(image)

        # Process the extracted text and add it to combined_data
        combined_data = []
        lines = extracted_text.split('\n')
        for i, line in enumerate(lines):
            if line.strip():  # Ignore empty lines
                # Using 'unknown' as the page number for image files
                metadata = f"{file_name} | 'text' | page 1 | line {i+1} | length {len(line)}"
                combined_data.append(('1', i+1, 'text', metadata, line))

        if combined_data:
            process_and_chunk_data(combined_data, file_name)
        # Writing the extracted data to the output file
        if not os.path.exists(os.path.dirname(output_path)):
            os.makedirs(os.path.dirname(output_path))

        with open(output_path, 'w', encoding='utf-8') as file:
            for data in combined_data:
                file.write(f"{data[3]}: {data[4]}\n")

        return combined_data

    except Exception as e:
        print(f"An error occurred while processing the image: {str(e)}")
        return []

# -----------------------------------------------------------------
def extract_data_from_docx(docx_file_path, output_path, file_name):
    """
    Extract data from a DOCX file.

    Args:
    - docx_file_path (str): Path to the DOCX file.
    - file_name (str): Name of the file being processed.

    Returns:
    - list: Extracted data combined from the DOCX file.
    """
    doc = Document(docx_file_path)
    combined_data = []

    # Process paragraphs
    page_number = 1  # DOCX files don't have page numbers; you'll need to handle this as needed.
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "":
            continue
        line_text = para.text.strip()
        metadata = f"{file_name} | 'text' | page {page_number} | para {i+1} | length {len(line_text)}"
        combined_data.append((page_number, i+1, 'text', metadata, line_text))

    # Process tables
    for table_index, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)

        # Now we create a DataFrame to handle the table as structured data
        df = pd.DataFrame(table_data)
        
        # Convert the DataFrame to a JSON string for easier processing later on
        table_content = df.to_json(orient='split')
        
        # Add the table data as a single item in combined_data
        table_metadata = f"{file_name} | 'table' | page {page_number} | table {table_index+1} | length {len(table_content)}"
        combined_data.append((page_number, 'N/A', 'table', table_metadata, table_content))

    if combined_data:
        process_and_chunk_data(combined_data, file_name)

    if not os.path.exists(os.path.dirname(output_path)):
        os.makedirs(os.path.dirname(output_path))

    with open(output_path, 'w', encoding='utf-8') as file:
        for data in combined_data:
            file.write(f"{data[3]}: {data[4]}\n")

    return combined_data
# ----------------------------------------------------
def write_to_combined_file(data, filepath, delimiter=None, save_as_csv=False):
    """
    Write combined data to a file.

    Args:
    - data (list): The data to write, consisting of tuples.
    - filepath (str): Path to the output file.
    - delimiter (str, optional): Delimiter to use between data elements. Defaults to '|'.
    - save_as_csv (bool, optional): Flag to save table data as CSV. Defaults to False.

    This function processes each item in the data list and writes it to the specified file. 
    It handles text and table data differently and can save table data as CSV if specified.
    """
    if delimiter is None:
        delimiter = DELIMITER_TBL  # Use default delimiter if none provided

    with open(filepath, 'w', encoding="utf-8") as f:
        for item in data:
            try:
                if item[2] == 'text':
                    # Writing text data to file
                    f.write(f"{item[0]}{delimiter}{item[1]}{delimiter}{item[2]}{delimiter}{item[3]}{delimiter}{item[4]}\n")
                elif item[2] == 'table':
                    # Processing table data
                    table_content = item[4]
                    if save_as_csv and isinstance(table_content, list):
                        df = pd.DataFrame.from_records(table_content)
                        table_as_csv = df.to_csv(index=False, header=False, sep=delimiter)
                        f.write(table_as_csv)
                    else:
                        f.write(str(table_content) + '\n')
                else:
                    log_message(f"Unexpected data type: {item}")
            except Exception as e:
                log_message(f"Error processing item: {e}")

    log_message(f"Writing {len(data)} items to {filepath}")

#-----------------------------------------------------------------------------

# VectorDB mapping: 'key' is the name shown to the user, 'value' is the storage path
VECTORDB_MAP = {
    'Default (In-memory)': None,  # In-memory client
    'ContractsDB': VECTORDB_DIR_SPACE_FOLDER  # Persistent client
}

def initialize_client(vector_db_name):
    """
    Initialize a ChromaDB client based on the provided VectorDB name.

    Args:
    - vector_db_name (str): The name of the VectorDB database.

    Returns:
    - chromadb.Client or chromadb.PersistentClient: An instance of a ChromaDB client.

    This function checks if the VectorDB name provided corresponds to a persistent storage path. 
    If so, it returns a PersistentClient; otherwise, it returns a standard in-memory Client.
    """
    persist_path = VECTORDB_MAP.get(vector_db_name)
    if persist_path:
        # If a persistent storage path is found, return a PersistentClient
        return chromadb.PersistentClient(path=persist_path)
    else:
        # If no persistent storage is associated, return an in-memory client
        return chromadb.Client()

#-----------------------------------------------------------------------------------
def ChromaDB_Initialize(client):
    """
    Initialize or retrieve a collection in ChromaDB.

    Args:
    - client (chromadb.Client): An instance of ChromaDB client.

    Returns:
    - Collection: A ChromaDB collection object.

    This function checks if a collection with the name specified in VECTORDB_NAME_SPACE 
    already exists within the provided ChromaDB client. If the collection exists, it retrieves 
    and returns this collection. If not, it creates a new collection with that name and then 
    returns it.
    """

    # Check if the collection specified by VECTORDB_NAME_SPACE already exists
    if VECTORDB_NAME_SPACE in [collection.name for collection in client.list_collections()]:
        # Notify user if the collection exists
        st.write(f"Collection {VECTORDB_NAME_SPACE} already exists.")
        # Retrieve or create the collection
        collection = client.get_or_create_collection(name=VECTORDB_NAME_SPACE)
    else:
        # Create the collection if it does not exist
        collection = client.get_or_create_collection(name=VECTORDB_NAME_SPACE)

    # Return the collection object
    return collection

#-----------------------------------------------------------------------------------
# Function to write a chunk of data to a file
def write_chunk_to_file(chunk, data_type, page_number, chunk_number, doc_id):
    filename = f"{data_type}_{doc_id}_Page_{page_number}_Chunk_{chunk_number}.txt"
    filepath = os.path.join(SMART_CHUNK_RESULT_OUTPUT_FOLDER, filename)
    
    with open(filepath, 'w', encoding='utf-8') as file:
        file.write(chunk)

#-----------------------------------------------------------------------------------
def flatten_data_structure_debug(data, doc_id):
    """
    Flatten and process the data structure for debugging purposes.

    Args:
    - data (list): A list of tuples containing extracted data from a document.
    - doc_id (str): The document identifier.

    The function iterates over the provided data list, where each item is expected 
    to be a tuple with 5 elements: page number, y-position, data type, information, 
    and content. It flattens this structure and reformats the information based 
    on specific criteria, adding metadata about the chunk number and size.
    """

    flattened_data = []

    for i, item in enumerate(data):
        # Expecting each item to be a tuple with 5 elements
        if len(item) != 5:
            # Skip items that do not match the expected structure
            continue
        
        # Unpack the tuple
        page_num, y_pos, data_type, info, content = item[:5]

        # Determine the chunk size based on the data type
        if data_type == 'text':
            chunk_size = len(content.split())  # Count words for text chunk size
        elif data_type == 'table':
            chunk_size = len(content)  # Assuming 'content' for tables is a list of lists representing rows
        
        # Replace placeholders in the info string with actual values
        metadata_info = info.replace('CHUNK_NUMBER_PLACEHOLDER', str(get_next_sequence())).replace('CHUNK_SIZE_PLACEHOLDER', str(chunk_size))

        # Add the processed data to the flattened list
        flattened_data.append((page_num, y_pos, data_type, metadata_info, content))
    
    return flattened_data

#-----------------------------------------------------------------------------------
def process_and_chunk_data(combined_data, doc_id):
    """
    Processes and chunks data extracted from a document.

    This function sorts and organizes the data by page number and type (table or text), 
    then processes it based on its type. Tables are processed and written as chunks directly, 
    while text data is further split into smaller chunks if needed.

    Args:
        combined_data (list of tuples): The data to be processed, typically extracted from a document.
        doc_id (str): The identifier for the document from which the data was extracted.

    Returns:
        None: The function does not return anything but writes processed data into files.

    Each tuple in the 'combined_data' list should have the following format:
        (Page Number, Y-position, Type, Info, Content)
    
    The function performs the following steps:
    1. Resets the sequence number for chunking.
    2. Flattens the combined data structure.
    3. Sorts and groups the data by page number and type.
    4. Processes each group:
       - If the data type is 'table', each table is converted into a string and written to a file as a chunk.
       - If the data type is 'text', the text is either written as a single chunk (if small enough) or split into smaller chunks.
    5. Writes each chunk to a file with appropriate metadata.
    """
    reset_sequence_number()
    
    # Flatten the combined_data structure
    flattened_data = flatten_data_structure_debug(combined_data, doc_id)

    #st.write(type(flattened_data), flattened_data[:5])

    # Create DataFrame from the flattened data
    df = pd.DataFrame(flattened_data, columns=['Page Number', 'Y-position', 'Type', 'Info', 'Content'])
    
    #st.write(df.columns)

    df_sorted = df.sort_values(by=['Page Number', 'Y-position'])
    
    # Group by 'Page Number' and 'Type'
    grouped = df_sorted.groupby(['Page Number', 'Type'])

    for (page_number, data_type), group in grouped:
        reset_sequence_number()
        log_message(f"Processing {data_type} on page {page_number} with {len(group)} items")

        if data_type == 'table':
                for index, table in enumerate(group.itertuples(), start=1):
                    chunk_number = get_next_sequence()
                    chunk_content = table.Content  # this is the list you need to convert to string
                    chunk_size = len(chunk_content)
                    metadata = f"{doc_id} | 'table' | {page_number} | {chunk_number} | {chunk_size}\n"

                    # Convert table list to a string representation (if necessary)
                    if isinstance(chunk_content, list):
                        # Example conversion: join rows with newline and columns with tabs
                        chunk_content_str = '\n'.join([' '.join(map(str, row)) for row in chunk_content])
                    else:
                        chunk_content_str = chunk_content

                    # concatenate the metadata with the chunk_content_str
                    full_content = metadata + chunk_content_str
                    write_chunk_to_file(full_content, 'table', page_number, chunk_number, doc_id)

        elif data_type == 'text':
            combined_text = ' '.join(group['Content'].astype(str)).strip()
            if len(combined_text) <= CHUNK_SIZE:
                chunk_number = get_next_sequence()
                metadata = f"{doc_id} | 'text' | {page_number} | {chunk_number} | {len(combined_text)}\n"
            
                write_chunk_to_file(metadata + combined_text, 'text', page_number, chunk_number, doc_id)
            else:
                
                text_chunks = text_splitter.split_text(combined_text)  # Assuming 'split_text' is the correct method
                for chunk_index, chunk in enumerate(text_chunks, start=1):
                    chunk_number = get_next_sequence()
                    metadata = f"{doc_id} | 'text' | {page_number} | {chunk_number} | {len(chunk)}\n"
                    write_chunk_to_file(metadata + chunk, 'text', page_number, chunk_number, doc_id)

    log_message("Finished processing and chunking of data.")
# -----------------------------------------------------------------------------------------
def safe_convert_to_dataframe(data):
    """
    Safely converts a dictionary or other data types to a Pandas DataFrame.

    This function handles the conversion of different data structures to a DataFrame. If the input is a dictionary with scalar values (not lists), it creates a single-row DataFrame. Otherwise, it converts the input using standard DataFrame conversion.

    Args:
        data (dict or any): The data to be converted into a DataFrame. It can be a dictionary or any data type that is compatible with Pandas DataFrame conversion.

    Returns:
        pandas.DataFrame: A DataFrame created from the input data.

    Note:
        - If the input is a dictionary with scalar values, each key-value pair becomes a column with the value as a single-row entry.
        - If the dictionary contains non-scalar values (like lists), or if the input is not a dictionary, the function proceeds with standard DataFrame conversion methods.
    """
    if isinstance(data, dict):
        # Check if the dictionary values are scalars (not lists)
        if all(isinstance(value, (int, float, str)) for value in data.values()):
            # Creating a DataFrame from a dictionary with scalar values, using an arbitrary index [0]
            return pd.DataFrame({k: [v] for k, v in data.items()})
        else:
            # If values are not scalar, convert as usual
            return pd.DataFrame(data)
    return data

#-----------------------------------------------------------------------------------
def get_table_count(pdf, strategy):
    """
    Counts the number of tables in a PDF document using a specific extraction strategy.

    This function iterates through each page of the given PDF document and uses the specified strategy to find tables. It sums up the total number of tables found across all pages.

    Args:
        pdf (pdfplumber.PDF): The PDF document object to be analyzed. The PDF should be opened using pdfplumber.
        strategy (dict): A dictionary defining the table extraction strategy. This should be compatible with pdfplumber's table finding settings.

    Returns:
        int: The total number of tables found in the PDF document using the given strategy.

    Example:
        strategy = {"vertical_strategy": "lines", "horizontal_strategy": "text"}
        pdf = pdfplumber.open('sample.pdf')
        table_count = get_table_count(pdf, strategy)
        pdf.close()
    """
    total_tables = 0
    for page in pdf.pages:
        tables = page.find_tables(table_settings=strategy)
        total_tables += len(tables)
    return total_tables

#-----------------------------------------------------------------------------------
def get_optimal_extraction_strategy(pdf, selected_strategy):
    """
    Determines the optimal table extraction strategy for a given PDF document.

    This function iterates through a set of predefined extraction strategies, applying each to the given PDF to count the number of tables extracted. It identifies the strategy that yields the highest table count. If two strategies result in the same count, it prioritizes the 'lines_lines' strategy.

    Args:
        pdf (pdfplumber.PDF): The PDF document object to be analyzed. The PDF should be opened using pdfplumber.
        selected_strategy (dict): A dictionary of the selected strategy to use for table extraction.

    Returns:
        tuple: A tuple containing the name of the optimal strategy and the maximum number of tables found using that strategy.

    Example:
        pdf = pdfplumber.open('sample.pdf')
        selected_strategy = {"vertical_strategy": "lines", "horizontal_strategy": "text"}
        optimal_strategy, max_tables = get_optimal_extraction_strategy(pdf, selected_strategy)
        pdf.close()
        print(f"Optimal strategy: {optimal_strategy}, Tables found: {max_tables}")

    Note:
        This function relies on the global EXTRACTION_STRATEGIES dictionary, which should be predefined with all possible strategies.
    """
    max_tables_found = 0
    optimal_strategy_name = None
    for strategy_name, strategy in EXTRACTION_STRATEGIES.items():
        current_table_count = get_table_count(pdf, strategy)
        log_message(f"{strategy_name} found {current_table_count} tables.")

        # Prioritize 'lines_lines' strategy if table count is the same
        if (current_table_count > max_tables_found) or (current_table_count == max_tables_found and strategy_name == 'lines_lines'):
            max_tables_found = current_table_count
            optimal_strategy_name = strategy_name

    log_message(f"Using {optimal_strategy_name} which found the most tables: {max_tables_found}")
    return optimal_strategy_name, max_tables_found

#-----------------------------------------------------------------------------------
def check_bboxes(word, table_bbox):
    """
    Determine if a word's bounding box is inside a given table's bounding box.

    Args:
        word (dict): Word dictionary with bbox details.
        table_bbox (tuple): Bounding box of the table.
    
    Returns:
        bool: True if the word is inside the table's bounding box, otherwise False.
    """
    l = word['x0'], word['top'], word['x1'], word['bottom']
    r = table_bbox

    return l[0] > r[0] and l[1] > r[1] and l[2] < r[2] and l[3] < r[3]

# ---------------------------------------------------------------------------------------------
def determine_optimal_strategy(pdf, selected_strategy):
    """
    Determines and returns the optimal table extraction strategy for a given PDF document, with an option to specify a preferred strategy.

    This function either utilizes a user-specified strategy or iterates through a set of predefined strategies to find the one that optimally extracts tables from the given PDF. If a preferred strategy is provided, it is used; otherwise, the function calls `get_optimal_extraction_strategy` to determine the best strategy based on the maximum table count.

    Args:
        pdf (pdfplumber.PDF): The PDF document object for which the strategy is to be determined.
        selected_strategy (dict, optional): The preferred strategy to use for table extraction. If not provided, the function will determine the optimal strategy.

    Returns:
        tuple: A tuple containing the name of the optimal strategy and the maximum number of tables found using that strategy.

    Example:
        pdf = pdfplumber.open('sample.pdf')
        selected_strategy = {"vertical_strategy": "lines", "horizontal_strategy": "text"}
        optimal_strategy, max_tables = determine_optimal_strategy(pdf, selected_strategy)
        pdf.close()
        print(f"Optimal strategy: {optimal_strategy}, Tables found: {max_tables}")

    Note:
        If 'selected_strategy' is None or not provided, the function will evaluate all predefined strategies in the EXTRACTION_STRATEGIES dictionary to find the most effective one.
    """
    # Logic to determine the optimal strategy
    if selected_strategy:
        return get_optimal_extraction_strategy(pdf, selected_strategy)
    else:
        return get_optimal_extraction_strategy(pdf)

# --------------------------------------------------------------------------------
def generate_and_save_table_chunk(page_number, y_position, content_type, metadata, processed_table_extract, doc_id):
    """
    Generates and saves a table data chunk to a file with corresponding metadata.

    This function processes a given table extract (in list format), converts it to a string representation, and then saves it as a chunk in a designated file. It also generates metadata for the chunk, including document ID, content type, page number, chunk number, and chunk size, which is prepended to the actual content.

    Args:
        page_number (int): The page number of the PDF document where the table is located.
        y_position (int): The Y-position of the table on the page (unused in current implementation).
        content_type (str): The type of content, usually 'table' for table data.
        metadata (str): Additional metadata about the chunk, typically including document source, page number, paragraph number, etc.
        processed_table_extract (list of list of str): The table content extracted from the PDF, where each inner list represents a row and each item a cell content.
        doc_id (str): The unique identifier for the document being processed.

    Example:
        # Assuming necessary variables and table content are already defined
        generate_and_save_table_chunk(1, 100, 'table', 'Sample Metadata', processed_table_extract, 'Doc123')

    Note:
        The function relies on 'write_chunk_to_file' for actual file writing. The 'y_position' argument is currently not utilized in the chunk metadata generation.
    """
    # Convert processed_table_extract to string representation
    table_content_as_string = '\n'.join([' '.join(row) for row in processed_table_extract])

    # Create chunk metadata
    chunk_number = get_next_sequence()
    chunk_size = len(table_content_as_string)
    chunk_metadata = f"{doc_id} | {content_type} | {page_number} | {chunk_number} | {chunk_size}\n"

    # Concatenate metadata with chunk content
    full_content = metadata + table_content_as_string

    # Save to file
    write_chunk_to_file(full_content, content_type, page_number, chunk_number, doc_id)

# -------------------------------------------

import re

def clean_content(content):
    """
    Cleans the given text content by removing characters not in the allowed list.

    This function takes a string, splits it into words, and removes any characters that are not in the defined set of allowed characters from each word. The allowed characters include letters (both uppercase and lowercase), numbers, and a specific set of symbols. The function then reassembles the cleaned words back into a single string.

    Args:
        content (str): The text content to be cleaned.

    Returns:
        str: The cleaned text content, where each word has been stripped of disallowed characters.

    Example:
        cleaned_text = clean_content("This is GenAI @ Interface.")
        print(cleaned_text)  # Output: "This is GenAI Interface."

    Note:
        The function returns a single space (' ') if the input content is empty or None. Also, it maintains the original spacing of words.
    """
    if not content:
        return ' '
    allowed_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789[]{}_-.,?!*&%$@()|/+\"':; "
    words = content.split()
    cleaned_words = [''.join([c if c in allowed_chars else '' for c in word]) for word in words]
    return ' '.join(cleaned_words)

#-----------------------------------------------------------------------------------

def process_table_row(row):
    """
    Cleans each cell in a table row by removing unwanted characters.

    This function iterates over each cell in the given row (which is expected to be a list of strings, with each string representing a cell). It applies the 'clean_content' function to each cell, effectively cleaning the cell's content by removing characters not in the allowed list.

    Args:
        row (list of str): A list of strings, where each string is the content of a cell in a table row.

    Returns:
        list: A list of cleaned strings representing the processed table row.

    Example:
        processed_row = process_table_row(["Hello, World!", "123 Main St."])
        print(processed_row)  # Output: ["Hello World", "123 Main St"]

    Note:
        If a cell is `None`, it is replaced with a single space (' ').
    """

    return [clean_content(cell) if cell is not None else ' ' for cell in row]
#-----------------------------------------------------------------------------------

import pdfplumber
from operator import itemgetter

def extract_and_order_data_from_pdf(page, optimal_strategy_name, doc_id):
    """
    Extracts and orders data from a single page of a PDF document.

    This function processes both textual and tabular data on a given page. It first extracts tables based on a predefined EXTRACTION_STRATEGY, then identifies non-table words and clusters them to form textual lines. Each table and text line is processed, cleaned, and added to a list with corresponding metadata.

    Args:
        page (pdfplumber.page.Page): The page object from pdfplumber.
        optimal_strategy_name (str): The name of the optimal strategy chosen for table extraction.
        doc_id (str): A unique identifier for the document being processed.

    Returns:
        list: A list of tuples, each containing page number, Y-position, data type ('text' or 'table'), metadata, and content (either text or table data).

    Example:
        items = extract_and_order_data_from_pdf(pdf_page, 'lines_text', 'doc123')
        for item in items:
            print(item)

    Note:
        The function uses a fixed EXTRACTION_STRATEGY for tables and a tolerance-based clustering approach for non-table words.
        Each item's metadata includes the document ID, content type, page number, line or table index, and the size of the content.
    """

    EXTRACTION_STRATEGY ={"vertical_strategy": "lines", 
                         "horizontal_strategy": "text", 
                          "snap_tolerance": 4}
    items = []
    #tables = page.find_tables(table_settings=EXTRACTION_STRATEGIES[optimal_strategy_name])
    tables = page.find_tables(table_settings=EXTRACTION_STRATEGY)
    non_table_words = [word for word in page.extract_words() if not any([check_bboxes(word, table.bbox) for table in tables])]
    clusters = pdfplumber.utils.cluster_objects(non_table_words, itemgetter('doctop'), tolerance=10)

    # Process text clusters

    for idx, cluster in enumerate(clusters):
        line_top = cluster[0]['doctop']
        line_text = ' '.join([word['text'] for word in cluster])
        line_text = clean_content(line_text).strip()
        metadata = f"{doc_id} | 'text' | {page.page_number} | {idx+1} | {len(cluster)}"
        items.append((page.page_number, line_top, 'text', metadata, line_text))

  
    # Process tables
    for idx, table in enumerate(tables):
        bbox = table.bbox
        table_extract = table.extract()
        processed_table_extract = [process_table_row(row) for row in table_extract]
        
        # Clean each cell in the table
        for i, row in enumerate(processed_table_extract):
            processed_table_extract[i] = [clean_content(cell) for cell in row]
        
        chunk_size = sum(len(' '.join(row)) for row in processed_table_extract)
        metadata = f"{doc_id} | 'table' | {page.page_number} | {idx+1} | {chunk_size}"
        
        # Replace tabs with spaces in the table content
        table_content_as_string = '\n'.join([' '.join(row) for row in processed_table_extract])
        
        items.append((page.page_number, bbox[1], 'table', metadata, table_content_as_string))

    return items
# ----------------------------------------------------------------------------------
import ahocorasick_rs

from ahocorasick_rs import AhoCorasick, MatchKind

def find_topics_in_text(content, topics):
    """
    Find occurrences of topics in the given content using Aho-Corasick algorithm.

    Args:
        content (str): Text content to search.
        topics (list): List of topics to search for.

    Returns:
        list: List of tuples containing the topic and its position in the content.
    """
    automaton = AhoCorasick(topics, matchkind=MatchKind.LeftmostLongest)

    try:
        # If content is not a string, try to convert it to string
        content = str(content)
    except Exception as e:
        error_message = f"Error converting content to string: {str(e)}"
        st.write(error_message)
        return [(-1, -1, -1, error_message)]  # Return a default value with an error message

    try:
        matches = automaton.find_matches_as_indexes(content)
    except Exception as e:
        error_message = f"Error in Aho-Corasick algorithm: {str(e)}"
        st.write(error_message)
        return [(-1, -1, -1, error_message)]  # Return a default value with an error message

    st.write("Matches:")
    if not matches:  # Check if no matches are found
        error_message = "No matches found."
        st.write(error_message)
        matches = [(-1, -1, -1, error_message)]  # Assign a default value with an error message

    flattened_matches = [(match[0], match[1], match[2]) for match in matches] if matches else [(-1, -1, -1, "Unknown error")]

    st.write(f"Flattened Matches: {flattened_matches}")  # Add this line for debugging

    return flattened_matches

# --------------------------------------------------------------------------------
# Function to process text chunks
def process_text_chunks(group, doc_id, page_number, CHUNK_SIZE, text_splitter, topics_to_find):
    combined_text = ' '.join(group['Content'].astype(str)).strip()

    if len(combined_text) <= CHUNK_SIZE:
        chunk_number = get_next_sequence()
        topic_matches = find_topics_in_text(combined_text, topics_to_find)
        topics_found = ', '.join([topic for topic, _, _ in topic_matches])
        metadata = f"{doc_id} | 'text' | {page_number} | {chunk_number} | {len(combined_text)} | Topics: {topics_found}\n"
        full_content = metadata + combined_text
        write_chunk_to_file(full_content, 'text', page_number, chunk_number, doc_id)
    else:
        text_chunks = text_splitter.split_text(combined_text)
        for chunk_index, chunk in enumerate(text_chunks, start=1):
            chunk_number = get_next_sequence()
            topic_matches = find_topics_in_text(chunk, topics_to_find)
            topics_found = ', '.join([topic for topic, _, _ in topic_matches])
            metadata = f"{doc_id} | 'text' | {page_number} | {chunk_number} | {len(chunk)} | Topics: {topics_found}\n"
            full_content = metadata + chunk
            write_chunk_to_file(full_content, 'text', page_number, chunk_number, doc_id)

def process_table_chunks(table, doc_id, page_number, topics_to_find):
    """
    Processes a single table chunk, finds topics in it, and writes it to a file.
    """
    chunk_number = get_next_sequence()

    # Extract the content from the table
    chunk_content = str(table['Content'])

    # Additional check for Pandas DataFrame row
    if isinstance(chunk_content, pd.Series):
        # Assuming the relevant text is in the '_1' column, update as needed
        chunk_content = str(chunk_content['_1'])

    # Find topics in the chunk
    try:
        topic_matches = find_topics_in_text(chunk_content, topics_to_find)
        topics_found = ', '.join([topic for topic, _, _ in topic_matches]) if topic_matches else 'NULL'
    except Exception as e:
        error_message = f"Error in find_topics_in_text: {str(e)}"
        st.write(error_message)
        return

    chunk_size = len(chunk_content)
    metadata = f"{doc_id} | 'table' | {page_number} | {chunk_number} | {chunk_size} | Topics: {topics_found}\n"
    full_content = metadata + chunk_content
    write_chunk_to_file(full_content, 'table', page_number, chunk_number, doc_id)
# ------------------------------------------------------------------------------------------------

import re
from urllib.parse import urlparse

# -------------------------------------------------------------------------------------

import re
from urllib.parse import urlparse

# Adjusted function with a broader regex pattern
def process_deep_links(text):
    """
    Identifies and processes deep links within the text content.
    """
    # Updated regex pattern to catch URLs that start with 'www.' and do not necessarily have 'http://' or 'https://'
    url_pattern = r'\b(?:https?://|www\.)[a-zA-Z0-9-\.]+\.[a-zA-Z]{2,}(?:\S*)\b'
    urls = re.findall(url_pattern, text)
    return [url for url in urls if urlparse(url).scheme in ['http', 'https'] or url.startswith('www.')]
#------------------------------------------------------------------------------

import requests
from bs4 import BeautifulSoup

def ensure_url_has_scheme(url):
    if not url.startswith(('http://', 'https://')):
        return 'https://' + url
    return url

def is_page_not_found(content):
    # Check for common phrases used in 404 error pages
    common_404_phrases = ["page not found", "404 error", "not found"]
    return any(phrase in content.lower() for phrase in common_404_phrases)

def process_url_with_beautifulsoup(url):
    """
    Processes a URL by making a request and using BeautifulSoup to parse the HTML content.
    It also checks if the page content indicates a 'Page Not Found' error.

    Args:
        url (str): A URL to be processed.

    Returns:
        str: The processed text content of the URL, or an error message if processing fails.
    """
    try:
        url_with_scheme = ensure_url_has_scheme(url)
        response = requests.get(url_with_scheme)
        soup = BeautifulSoup(response.content, 'html.parser')
        text_content = ' '.join(soup.stripped_strings)

        if response.status_code == 200 and not is_page_not_found(text_content):
            st.write(f'response code : {response.status_code} and url {url_with_scheme}')
            return text_content
        elif is_page_not_found(text_content):
            log_message (f"Page not found (404 Error) at {url_with_scheme}")
        else:
            log_message (f"Error: URL returned status code {response.status_code}")

    except requests.RequestException as e:
        log_message (f"Error processing URL {url}: {e}")
# ----------------------------------------------------------------------------
def ChromaDB_Process_Add_Data_url(client, collection, coimbined_data):
    GET_FILE_NAMES = glob.glob(f"{SMART_CHUNK_RESULT_OUTPUT_FOLDER}/*url*.txt")
    
    if not GET_FILE_NAMES:
        st.write("No .txt file found in the input folder\n")
        return None, collection

    # Capture the current date-time now, to be used for all files processed in this run
    current_load_date = datetime.now().isoformat()

    for file_name in GET_FILE_NAMES:
        with open(file_name, 'r', encoding='utf-8') as text_file:
            # Read the first line for custom metadata and split it
            metadata_firstline = text_file.readline().rstrip().split('|')

            # Ensure the split metadata has the expected number of elements
            if len(metadata_firstline) >= 4:
                doc_id = metadata_firstline[0].strip()
                data_type = metadata_firstline[1].strip()
                page_number = metadata_firstline[2].strip()
                chunk_number = metadata_firstline[3].strip()
                url = metadata_firstline[5].strip()

                # Read the remaining content
                remaining_content = text_file.read()

                # Prepare metadata including contract_id
                metadata = {
                    "source": doc_id,
                    "data_type": data_type,
                    "page_number": page_number,
                    "chunk_number": chunk_number,
		            "url_link" : url,
                    "load_date": current_load_date, 
                    "contract_id": 'Not Applicable',
                    "customer_id": 'Not Applicable',
                    "amendment_id": 'Not Applicable',
                    "validity_date": 'Not Applicable'  
                }

                # Generate content hash for unique document ID
                content_hash = generate_content_hash(remaining_content)

                # Add document to collection
                collection.add(documents=[remaining_content], metadatas=[metadata], ids=[content_hash])

            else:
                st.write(f"Invalid metadata format in file {file_name}")
                continue
        
    return doc_id, collection
# -------------------------------------------------------------------------------------
def write_chunk_to_url_file(chunk, data_type, page_number, chunk_number, doc_id):
    try:
        # Create the filename using the provided parameters
        filename = f"{data_type}_{doc_id}_Page_{page_number}_Chunk_{chunk_number}.txt"
        filepath = os.path.join(SMART_CHUNK_RESULT_OUTPUT_FOLDER, filename)
        with open(filepath, 'w', encoding='utf-8') as file:
            file.write(chunk)
        return True
    except Exception as e:
        # Log error or handle it as needed
        print(f"Error writing chunk url to file: {e}")
        return False
# --------------------------------------------------------------------------------------
# Function to process URL chunks
def process_url_chunks(content, doc_id, page_number, text_splitter, url=None):
    try:
        combined_text = ' '.join(content['Content'].astype(str)).strip()
        if len(combined_text) <= CHUNK_SIZE:
            chunk_number = get_next_sequence()
            metadata = f"{doc_id} | 'url' | {page_number} | {chunk_number} | {len(combined_text)} | {url if url else ''}\n"
            return write_chunk_to_url_file(metadata + combined_text, 'url', page_number, chunk_number, doc_id)
        else:
            success = True
            text_chunks = text_splitter.split_text(combined_text)
            for chunk_index, chunk in enumerate(text_chunks, start=1):
                chunk_number = get_next_sequence()
                metadata = f"{doc_id} | 'url' | {page_number} | {chunk_number} | {len(chunk)} | {url if url else ''}\n"
                if not write_chunk_to_url_file(metadata + chunk, 'url', page_number, chunk_number, doc_id):
                    success = False
            return success
    except Exception as e:
        # Log error or handle it as needed
        print(f"Error processing URL chunks: {e}")
        return False
# -----------------------------------------------------------------------------
import pandas as pd
import pdfplumber

def working_extract_data_from_pdf(INPUT_PDF_FILE, COMBINED_DATA_EXTRACT_RESULT_OUTPUT_PATH, doc_id, selected_strategy):
    """
    Extracts and processes data from a PDF file, combining and organizing the extracted data.

    This function opens a PDF file and processes each page to extract text and table content. It organizes this content into a DataFrame, sorts it, merges tables where applicable, and then processes each text and table chunk. The processed data is then written to a JSON file.

    Args:
        INPUT_PDF_FILE (str): The path to the input PDF file.
        COMBINED_DATA_EXTRACT_RESULT_OUTPUT_PATH (str): The path where the combined JSON output will be saved.
        doc_id (str): A unique identifier for the document being processed.
        selected_strategy (str): The name of the strategy to be used for data extraction.

    Returns:
        str: A JSON string of the combined and processed data.

    Example:
        working_extract_data_from_pdf('input.pdf', 'output.json', 'doc123', 'lines_lines')

    Note:
        The function uses the `extract_and_order_data_from_pdf` function to extract data from each page and `pdfplumber` for opening and reading the PDF file.
        It merges tables if they span across pages and processes each text and table chunk separately.
        The output is saved as a JSON file and also returned as a JSON string.
    """
    combined_data = []

    with pdfplumber.open(INPUT_PDF_FILE) as pdf:
        for page in pdf.pages:
            items = extract_and_order_data_from_pdf(page,selected_strategy, doc_id)
            combined_data.extend(items)

    df = pd.DataFrame(combined_data, columns=['Page Number', 'Y-position', 'Type', 'Info', 'Content'])
    df_sorted = df.sort_values(by=['Page Number', 'Y-position'])

    # Merging tables based on conditions
    prev_page_number = -1
    prev_y_position = float('inf')
    prev_index = None

    for index, row in df_sorted.iterrows():
        if row['Type'] == 'table':
            if row['Page Number'] == prev_page_number + 1 and row['Y-position'] < prev_y_position and prev_index is not None:
                df_sorted.at[prev_index, 'Content'] += '\n' + row['Content']
                df_sorted = df_sorted.drop(index)
            prev_index = index
            prev_page_number = row['Page Number']
            prev_y_position = row['Y-position']

    df_sorted = df_sorted.reset_index(drop=True)

    # Group and process each chunk
    grouped = df_sorted.groupby(['Page Number', 'Type'])
    for (page_number, data_type), group in grouped:
        reset_sequence_number()
        log_message(f"Processing {data_type} on page {page_number} with {len(group)} items")
        if data_type == 'table':
            for index, table in enumerate(group.itertuples(), start=1):
                topic_matches = find_topics_in_text(table, topics_to_find)
                process_table_chunks(table, doc_id, page_number, topic_matches)
        elif data_type == 'text':
            process_text_chunk(group, doc_id, page_number, CHUNK_SIZE, text_splitter, topics_to_find)
            # After extracting data from PDF
            combined_text = ' '.join(df['Content'].astype(str)).strip()
            if combined_text:
                #st.write(combined_text)
                urls = process_deep_links(combined_text)
    
    if urls:
        
        for url in urls:
            content = process_url_with_beautifulsoup(url)
            if isinstance(content, str):
                # Convert the content to a DataFrame and process as a text chunk
                st.write(clean_content(content))
                process_url_chunks(pd.DataFrame({'Content': [clean_content(content)]}), doc_id, page_number, CHUNK_SIZE, text_splitter, url)
                
    combined_json_data = df_sorted.to_json(orient='records', indent=4)
    with open(COMBINED_DATA_EXTRACT_RESULT_OUTPUT_PATH, 'w') as json_file:
        json_file.write(combined_json_data)

    return combined_json_data

# ------------- Chroma DB Initialization -------------------------------------------
def ChromaDB_Initialize(client):
    """
    Initializes or retrieves a collection within a ChromaDB client instance.

    This function checks if a specified collection (defined by the constant VECTORDB_NAME_SPACE) exists within the given ChromaDB client. If it exists, the existing collection is retrieved, otherwise, a new collection is created with the given name.

    Args:
        client (chromadb.Client or chromadb.PersistentClient): The ChromaDB client instance to initialize the collection in.

    Returns:
        chromadb.Collection: The initialized or retrieved collection.

    Example:
        client = chromadb.Client()
        collection = ChromaDB_Initialize(client)

    Note:
        - If VECTORDB_NAME_SPACE is already present in the client's list of collections, the existing collection is returned.
        - If VECTORDB_NAME_SPACE is not found, a new collection is created with that name and returned.
        - The function includes commented-out code for alternative behaviors like deleting an existing collection before creating a new one.
    """
    if VECTORDB_NAME_SPACE in [collection.name for collection in client.list_collections()]:
        st.write(f"Collection {VECTORDB_NAME_SPACE} already exists.")
        collection = client.get_or_create_collection(name=VECTORDB_NAME_SPACE)
    else:
        collection = client.get_or_create_collection(name=VECTORDB_NAME_SPACE)
    return collection
# -------------------------------------------------------------------------------------

def generate_content_hash(content):
    """
    Generates a SHA-256 hash for the given content.

    This function takes a string input and returns its SHA-256 hash. The content is first encoded to UTF-8 before hashing. SHA-256 is a cryptographic hash function that produces a fixed size (256-bit) hash. It is commonly used for security applications and integrity verification.

    Args:
        content (str): The content to be hashed.

    Returns:
        str: The SHA-256 hash of the provided content.

    Example:
        hashed_content = generate_content_hash("example content")

    Note:
        - SHA-256 hashes are widely used for verifying data integrity and in cryptographic systems.
        - The function requires the content to be a string. If working with other data types, convert them to a string format before passing to this function.
    """
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

# -----------------------------------------------------------------------------------------
import re
import json
from datetime import datetime

# Function to convert date format from DD-MM-YYYY to MM/DD/YYYY

def convert_date_format(date_str):
    """
    Converts a date string from 'DD-MonthFullName-YYYY' format to 'MM/DD/YYYY' format.

    This function takes a date string in the format of 'DD-MonthFullName-YYYY' (e.g., '22-October-2023') and converts it to the format 'MM/DD/YYYY' (e.g., '10/22/2023'). It uses Python's datetime library for parsing and formatting the date.

    Args:
        date_str (str): The date string in 'DD-MonthFullName-YYYY' format.

    Returns:
        str: The date string converted to 'MM/DD/YYYY' format.

    Raises:
        ValueError: If the input date_str is not in the expected format.

    Example:
        formatted_date = convert_date_format("22-October-2023")
        # Output: "10/22/2023"

    Note:
        - The function strictly expects the input date string to follow the 'DD-MonthFullName-YYYY' format.
        - A ValueError is raised if the input string does not match the expected format.
    """
    # Parse the date in DD-MonthFullName-YYYY format (e.g., 22-October-2023)
    date_obj = datetime.strptime(date_str, '%d-%B-%Y')

    # Convert to MM/DD/YYYY format
    new_format_date_str = date_obj.strftime('%m/%d/%Y')
    return new_format_date_str
# -----------------------------------------------------------------------------------------

# Define IDs regex patterns
amendment_pattern = re.compile(r"Amendment ID:\s*(\S+)")
contract_pattern = re.compile(r"Contract ID:\s*(\S+)")
customer_pattern = re.compile(r"([A-Z0-9-\.]+) \(Customer\)")
# Regex pattern for Validity Date
validity_date_pattern = re.compile(r"Valid if signed and submitted to Verizon by (\d{2}-\w+-\d{4})")

# extract IDs from Combined data
# ----------------------------------------------------------------------------
def extract_ids_from_combined_data(combined_data):
    """
    Extracts amendment ID, contract ID, customer ID, and validity date from combined data.

    This function processes the combined data, which is either a JSON string or a list of dictionaries, and extracts four specific pieces of information: amendment ID, contract ID, customer ID, and the validity date. Each of these items is searched within the 'Content' key of the dictionaries. If an item is found, its value is extracted and returned. The validity date is also formatted to 'MM/DD/YYYY' format.

    Args:
        combined_data (str or list): Combined data in JSON string format or as a list of dictionaries.

    Returns:
        tuple: A tuple containing four strings: amendment_id, contract_id, customer_id, and validity_date. If any of these items are not found, 'Unknown' is returned in their place.

    Note:
        - The function stops processing as soon as all four items are found.
        - If the combined data is a string, it attempts to parse it as JSON. If parsing fails, it returns 'Unknown' for all four items.
    """
    # Check if combined_data is a string and needs parsing
    if isinstance(combined_data, str):
        try:
            combined_data = json.loads(combined_data)
        except json.JSONDecodeError:
            st.write("Error decoding JSON")
            return "Unknown", "Unknown", "Unknown", "Unknown"

    amendment_id, contract_id, customer_id, validity_date = "Unknown", "Unknown", "Unknown", "Unknown"

    for item in combined_data:
        # Assuming item is a dictionary and has a 'Content' key
        content = item.get("Content", "")

        # Extracting amendment, contract, and customer IDs
        if amendment_id == "Unknown":
            amendment_match = amendment_pattern.search(content)
            if amendment_match:
                amendment_id = amendment_match.group(1)

        if contract_id == "Unknown":
            contract_match = contract_pattern.search(content)
            if contract_match:
                contract_id = contract_match.group(1)

        if customer_id == "Unknown":
            customer_match = customer_pattern.search(content)
            if customer_match:
                customer_id = customer_match.group(1)

        # Extracting and converting validity date
        if validity_date == "Unknown":
            validity_date_match = validity_date_pattern.search(content)
            if validity_date_match:
                submitted_date_str = validity_date_match.group(1)
                validity_date = convert_date_format(submitted_date_str)  # Convert date format

        if all(id != "Unknown" for id in [amendment_id, contract_id, customer_id, validity_date]):
            break

    return amendment_id, contract_id, customer_id, validity_date
# ----------------------------------------------------------------------------

# ---------------------------------------------------------------------------
def ChromaDB_Process_Add_Data_pdf(client, collection, coimbined_data):
    GET_FILE_NAMES = glob.glob(f"{SMART_CHUNK_RESULT_OUTPUT_FOLDER}/*.txt")
    
    if not GET_FILE_NAMES:
        st.write("No .txt file found in the input folder")
        return None, collection

    amendment_id, contract_id, customer_id, validity_date = extract_ids_from_combined_data(coimbined_data)

    # Capture the current date-time now, to be used for all files processed in this run
    current_load_date = datetime.now().isoformat()

    st.write("Amendment ID:", amendment_id)
    st.write("Contract ID:", contract_id)
    st.write("Customer ID:", customer_id)
    st.write("Validity Date:", validity_date)  # Displaying the validity date


    for file_name in GET_FILE_NAMES:
        with open(file_name, 'r', encoding='utf-8') as text_file:
            metadata_firstline = text_file.readline().rstrip().split('|')

            if len(metadata_firstline) >= 4:
                doc_id, data_type, page_number, chunk_number = metadata_firstline[0].strip(), metadata_firstline[1].strip(), metadata_firstline[2].strip(), metadata_firstline[3].strip()
                remaining_content = text_file.read()


                topics_found = ', '.join([topic for topic, _, _ in topics_to_find])

                metadata = {
                    "source": doc_id,
                    "data_type": data_type,
                    "page_number": page_number,
                    "chunk_number": chunk_number,
                    "load_date": current_load_date,  # Use the same load_date for all files
                    "contract_id": 'Not Applicable',
                    "customer_id": 'Not Applicable',
                    "amendment_id": 'Not Applicable',
                    "validity_date": 'Not Applicable',
                    "topics_found": topics_found  # Adding topics_found to metadata
                }

                content_hash = generate_content_hash(remaining_content)
                collection.add(documents=[remaining_content], metadatas=[metadata], ids=[content_hash])
            else:
                st.write(f"Invalid metadata format in file {file_name}")

    return True, collection
# ---------------------------------------------------------------------

def ChromaDB_Process_Add_Data_docx_img(client, collection, coimbined_data):
    GET_FILE_NAMES = glob.glob(f"{SMART_CHUNK_RESULT_OUTPUT_FOLDER}/*.txt")
    
    if not GET_FILE_NAMES:
        st.write("No .txt file found in the input folder\n")
        return None, collection

    # Capture the current date-time now, to be used for all files processed in this run
    current_load_date = datetime.now().isoformat()

    for file_name in GET_FILE_NAMES:
        with open(file_name, 'r', encoding='utf-8') as text_file:
            # Read the first line for custom metadata and split it
            metadata_firstline = text_file.readline().rstrip().split('|')

            # Ensure the split metadata has the expected number of elements
            if len(metadata_firstline) >= 4:
                doc_id = metadata_firstline[0].strip()
                data_type = metadata_firstline[1].strip()
                page_number = metadata_firstline[2].strip()
                chunk_number = metadata_firstline[3].strip()

                # Read the remaining content
                remaining_content = text_file.read()

                # Prepare metadata including contract_id
                metadata = {
                    "source": doc_id,
                    "data_type": data_type,
                    "page_number": page_number,
                    "chunk_number": chunk_number,
                    "load_date": current_load_date, 
                    "contract_id": 'Not Applicable',
                    "customer_id": 'Not Applicable',
                    "amendment_id": 'Not Applicable',
                    "validity_date": 'Not Applicable'  
                }

                # Generate content hash for unique document ID
                content_hash = generate_content_hash(remaining_content)

                # Add document to collection
                collection.add(documents=[remaining_content], metadatas=[metadata], ids=[content_hash])

            else:
                st.write(f"Invalid metadata format in file {file_name}")
                continue
        
    return doc_id, collection
# ----------------------------------------------------------------------------------------

def query_collection(collection, query, n_results, customer_id=None, additional_where=None, additional_where_document=None, source=None, contract_id=None, amendment_id=None, validity_date=None, get_latest=True, start_date=None, end_date=None):
    """
    Executes a query on a ChromaDB collection with various filtering options.

    This function allows querying a ChromaDB collection with specified criteria such as customer ID, source, contract ID, amendment ID, validity date, and additional custom filters. It supports both fetching the latest result or filtering results within a specific date range.

    Args:
        collection: ChromaDB collection object.
        query: Query text for the search.
        n_results: Number of results to fetch.
        customer_id: Filter for specific customer ID(s).
        additional_where: Additional custom filters for querying.
        additional_where_document: Additional document-level filters.
        source: Filter for the specific source.
        contract_id: Filter for specific contract ID.
        amendment_id: Filter for specific amendment ID.
        validity_date: Filter for specific validity date.
        get_latest: Flag to get only the latest result (default: True).
        start_date: Start date for filtering results (used when get_latest is False).
        end_date: End date for filtering results (used when get_latest is False).

    Returns:
        dict or list: Latest result as a dictionary if get_latest is True, or a list of filtered results within the specified date range.
        None: If no results are found or in case of an error.

    Raises:
        Streamlit error: If an error occurs during query execution.

    Notes:
        - The function constructs a where clause based on the provided filters.
        - If get_latest is True, the function returns the most recent result based on the 'load_date' metadata field.
        - If get_latest is False, the function returns all results within the specified date range.
        - The function logs the input query and the error message in case of an exception.
    """

    try:
        st.write(f"Input Query: {query}")

        # Prepare the query parameters
        query_params = {
            'query_texts': [query],
            'n_results': n_results
        }

        # Construct the where clause
        where_clause = {}

        # Adding filters
        if customer_id:
            customer_id_list = [customer_id] if not isinstance(customer_id, list) else customer_id
            where_clause['customer_id'] = {"$in": customer_id_list}
        if source:
            where_clause['source'] = {"$eq": source}
        if contract_id:
            where_clause['contract_id'] = {"$eq": contract_id}
        if amendment_id:
            where_clause['amendment_id'] = {"$eq": amendment_id}
        if validity_date:
            where_clause['validity_date'] = {"$eq": validity_date}
        if additional_where:
            where_clause.update(additional_where)

        if additional_where_document:
            query_params['where_document'] = additional_where_document

        if where_clause:
            query_params['where'] = where_clause

        # Query execution
        raw_results = collection.query(**query_params)

        if raw_results and "ids" in raw_results and "metadatas" in raw_results and "documents" in raw_results:
            results = [{"id": id, "metadata": metadata, "document": document} for id, metadata, document in zip(raw_results["ids"][0], raw_results["metadatas"][0], raw_results["documents"][0])]

            if get_latest:
                latest_result = sorted(results, key=lambda x: x['metadata'].get('load_date', datetime.min), reverse=True)[0]

                topics_found = ', '.join([topic for topic, _, _ in topics_to_find])

                latest_result['metadata']['topics_found'] = topics_found

                return latest_result
            else:
                # Filter results within the specified date range
                filtered_results = [result for result in results if start_date <= datetime.fromisoformat(result['metadata'].get('load_date', datetime.min.isoformat())) <= end_date]
                topics_found = ', '.join([topic for topic, _, _ in topics_to_find])
                for result in filtered_results:
                    topic_matches = find_topics_in_text(result['document'], topics_to_find)
                    topics_found = ', '.join([topic for topic, _, _ in topic_matches])
                    result['metadata']['topics_found'] = topics_found

                return filtered_results
        else:
            st.write("No results found or unexpected format of query results.")
            return None

    except Exception as e:
        st.error(f"Error executing query: {str(e)}")
        return None

#----------------------------------------------------------------------------------------
# Function for simulated chat input
def user_input(user_question):
    response = st.session_state.conversation({'question': user_question})
    st.session_state.chatHistory = response['chat_history']
    for i, message in enumerate(st.session_state.chatHistory):
        if i % 2 == 0:
            st.write("Human: ", message.content)
        else:
            st.write("Bot: ", message.content)
# ----------------------------------------------------------------------
# Function to remove duplicates from results based on document content
def remove_duplicates_from_results(results):
    """
    Removes duplicate entries from a list of search results based on document content.

    This function iterates through a list of results, each either a dictionary containing document data or a string representing the document content. It utilizes a SHA-256 hash to identify and remove duplicates, ensuring each document in the returned list is unique.

    Args:
        results (list): A list of results, where each result is either a dictionary with a 'document' key or a string representing the document content.

    Returns:
        list: A list of unique results with duplicates removed.

    Note:
        - The uniqueness of documents is determined by the hash of their content.
        - If a result is a dictionary, the function uses the value of the 'document' key for duplicate checking. Otherwise, the result is treated as a string.
        - The function maintains the original order of the results, with duplicates removed.
    """
    unique_results = []
    seen_documents = set()

    for result in results:
        # Check if the result is a dictionary
        if isinstance(result, dict):
            doc_content = result.get('document', '')
        else:
            # If result is not a dictionary, treat it as a string
            doc_content = result

        # Generate a hash of the document content
        content_hash = hashlib.sha256(doc_content.encode('utf-8')).hexdigest()

        if content_hash not in seen_documents:
            seen_documents.add(content_hash)
            unique_results.append(result)

    return unique_results
# ---------------------------------------------------------------------------------
import openllm
# --------------------------------------------------------------------------------------

@st.cache(allow_output_mutation=True)
def cached_load_llm(model_id, use_openllm, model_dir):
    """
    Loads a Language Model (LLM) from a specified directory. The function supports both OpenLLM and local Transformers models.

    This function is designed to cache the loaded model for faster subsequent accesses. It uses Streamlit's caching mechanism to store the model in memory, reducing load times on subsequent calls with the same arguments.

    Args:
        model_id (str): Identifier of the model to load.
        use_openllm (bool): Flag indicating whether to load an OpenLLM model. If False, a Transformers model will be loaded.
        model_dir (str): Directory path where the model is stored.

    Returns:
        object: A model pipeline object ready for text generation. This could be either an instance of openllm.LLM or a Transformers pipeline.

    Note:
        - The function assumes the necessary model files are present in the specified directory.
        - The `@st.cache` decorator is used to cache the function's output. It allows output mutation but ensures the same model object is returned for the same inputs.
        - The Transformers pipeline is configured for text generation with specific parameters like `max_length`, `do_sample`, `top_k`, etc.
    """
    if use_openllm:
        # Load OpenLLM model locally
        return openllm.LLM(model_dir)
    else:
        # Load local Transformers model
        model = AutoModelForCausalLM.from_pretrained(model_dir)
        tokenizer = AutoTokenizer.from_pretrained(model_dir)
        return pipeline("text-generation", model=model, tokenizer=tokenizer,
                        use_cache=True, device_map="auto",
                        max_length=512, do_sample=True, top_k=5, num_return_sequences=1)

# --------------------------------------------------------------------------------------------

import threading

def load_llm_async(model_id, use_openllm, callback):
    """
    Asynchronously loads a Language Learning Model (LLM) and triggers a callback function once the model is loaded. 
    The function supports loading both OpenLLM and Transformers models.

    This function is designed to load the model in a separate thread to prevent blocking the main execution flow, 
    particularly useful in asynchronous environments or web applications like Streamlit.

    Args:
        model_id (str): Identifier of the model to be loaded.
        use_openllm (bool): Flag indicating whether to load an OpenLLM model. If False, a Transformers model will be loaded.
        callback (function): A callback function to be invoked once the model is loaded. The loaded model is passed as an argument to this function.

    Returns:
        None: The function does not return a value but instead calls the provided callback function with the loaded model.

    Note:
        - The function uses `threading.Thread` to handle asynchronous loading.
        - It leverages the `cached_load_llm` function for loading and caching the Transformers model.
        - The model path is constructed based on the `MODEL_FOLDER` global variable and the `model_id`.
        - Streamlit's output (e.g., `st.write`) is used to display information about the model path.
    """
    model_path = os.path.join(MODEL_FOLDER, model_id)
    st.write(f'model_path : {model_path}')
    if use_openllm:
        # Asynchronously load OpenLLM model
        threading.Thread(target=lambda: callback(openllm.LLM(model_path))).start()
    else:
        # Asynchronously load Transformers model
        model = cached_load_llm(model_id, False, model_path)
        callback(model)

# -----------------------------------------------------------------------

def load_llm_callback(model):
    """
    Callback function for handling the loaded Language Learning Model (LLM). 
    It stores the loaded model in Streamlit's session state and triggers a rerun of the Streamlit app.

    This function is typically used as a callback in asynchronous model loading scenarios, 
    where the model is loaded in a separate thread and needs to be integrated into the Streamlit session once ready.

    Args:
        model (LLM object): The loaded Language Learning Model (LLM) object. 
                            This can be either an OpenLLM model or a Transformers model, depending on the loading process.

    Returns:
        None: The function does not return a value but updates the Streamlit session state and triggers a rerun of the app.

    Note:
        - The function assumes that Streamlit's session state is being used to manage application state.
        - It uses `st.experimental_rerun()` to refresh the Streamlit app, which is useful to update the UI with the loaded model.
        - The model is stored in the session state under the key 'loaded_model', allowing it to be accessed throughout the app.
    """
    st.session_state['loaded_model'] = model
    st.experimental_rerun()

# ----------------------------------------------------------------------------------------
# Function to process query with LLM
def process_query_with_llm(query, context, model_id, use_openllm):
    """
    Processes a query using a Language Learning Model (LLM) and returns the model's response. 
    The function uses a pre-loaded model stored in Streamlit's session state and supports both OpenLLM and Transformers models.

    Args:
        query (str): The query string to be processed by the LLM.
        context (str): The context string providing additional information relevant to the query.
        model_id (str): Identifier for the model. Used for logging or tracking purposes.
        use_openllm (bool): Flag indicating whether to use OpenLLM or a Transformers model.

    Returns:
        response (str or dict): The response from the LLM. The format can vary depending on the model used.
                                For OpenLLM, it's typically a string, while for Transformers, it might be a dictionary.

    Example:
        - If `use_openllm` is True, the query is directly passed to the OpenLLM model.
        - If `use_openllm` is False, a prompt combining both context and query is constructed and passed to the Transformers model.

    Note:
        - Assumes that the LLM has been previously loaded and stored in Streamlit's session state under the key 'loaded_model'.
        - It's important to ensure that the model is loaded and available in the session state before calling this function.
        - The function does not handle the loading of models. This should be handled separately.
    """
    llm = st.session_state.get('loaded_model')
    if use_openllm:
        response = llm(query)  # Query the OpenLLM model
    else:
        prompt = f"Context: {context}\nQuestion: {query}\nAnswer:"
        response = llm(prompt)  # Query the Transformers model
    return response
# ----------------------------------------------------------------------------------------

# JavaScript to copy text to clipboard
copy_js = """
<script>
function copyToClipboard(text) {
  const elem = document.createElement('textarea');
  elem.value = text;
  document.body.appendChild(elem);
  elem.select();
  document.execCommand('copy');
  document.body.removeChild(elem);
}
</script>
"""
# Streamlit component to call the JavaScript function
def copy_button(text, label='Copy the Result'):
    components.html(f"{copy_js}<button onclick='copyToClipboard(`{text}`)'>{label}</button>")

import base64

def download_link(object_to_download, download_filename, download_link_text):
    """
    Generates a link to download the given object_to_download.

    Args:
    - object_to_download (str or pd.DataFrame): The object to be downloaded.
    - download_filename (str): Filename and extension of file. e.g. mydata.csv, some_txt_output.txt
    - download_link_text (str): Text to display for download link.

    Returns:
    - str: HTML anchor tag to download object_to_download
    """
    if isinstance(object_to_download, str):
        # Strings are encoded to bytes and then base64 encoded
        b64 = base64.b64encode(object_to_download.encode()).decode()
    else:
        # Assume it's a DataFrame for now, but could be extended for other types
        b64 = base64.b64encode(object_to_download.to_csv(index=False).encode()).decode()

    return f'<a href="data:file/txt;base64,{b64}" download="{download_filename}">{download_link_text}</a>'

# ----------------------------------------------------------------------------------------------------------
def create_html_table(data):
    html = "<table style='border-collapse: collapse; font-size: small;'>"
    html += "<thead><tr>"
    headers = ['Metadata', 'Document', 'Distance']
    for header in headers:
        html += f"<th style='border: 1px solid black; padding: 8px;'>{header}</th>"
    html += "</tr></thead>"
    html += "<tbody>"

    for row_index, row in enumerate(data):
        html += "<tr>"
        for item in row:
            html += f"<td style='border: 1px solid black; padding: 8px;'>{item}</td>"
        html += "</tr>"

        # Add a new row for actions (Download and Copy)
        document_content = row[1]  # Assuming the document content is the second element in the row
        download_button = download_link(document_content, f"document_{row_index}.txt", "Download")
        copy_button = f"<button onclick='navigator.clipboard.writeText(\"{document_content}\")'>Copy Result</button>"
        html += f"<td style='border: 1px solid black; padding: 8px;' colspan='3'>{download_button} {copy_button}</td>"

    html += "</tbody></table>"
    return html

#---------------------------------------------------------------------------------
def display_results1(results, unique_results, max_display=NUMBER_SEARCH_QUERY):
    max_display = max(1, max_display)  # Ensures that max_display is at least 1

    try:
        if not results or 'ids' not in results:
            st.write("No results to display.")
            return

        table_data = []

        # Adjusted to handle the structure of 'unique_results'
        for idx, result in enumerate(unique_results[:max_display]):
            # Check if 'result' is a dictionary and contains necessary keys
            if isinstance(result, dict) and 'document' in result and 'metadata' in result:
                document = escape(result['document'])
                metadata = result['metadata']
                file_name = metadata.get('source', '')

                # Extract DOC_ID and other details from file_name
                match = re.match(r"(text_|table_)(.*?)_Page_(\d+)_Chunk_(\d+)_(.*)", file_name)
                if match:
                    _, DOC_ID, page_info, chunk_id, _ = match.groups()
                    fully_qualified_file_name = os.path.join(ARCHIVE_FOLDER, DOC_ID, 'SMART_CHUNK_RESULT_OUTPUT_FOLDER', f"{file_name}.txt")
                    file_link = f"<a href='file://{quote(fully_qualified_file_name)}' target='_blank'>{DOC_ID}</a>"
                    metadata_display = f"File Name: {file_link}<br>Page Number: {page_info}<br>Chunk #: {chunk_id}"
                    table_data.append([metadata_display, document, "N/A"])  # Distance is not available

        # Create and display the HTML table
        st.write(results)
        html_table = create_html_table(table_data)
        st.markdown(html_table, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

# ------------------------------------------------------------------------------------
def reset_database(client, db_directory):
    try:
        # Reset the ChromaDB database
        client.reset()
        st.sidebar.write("Database has been reset.")

        # Check if the directory exists
        if os.path.exists(db_directory):
            # Delete the directory and all its contents
            shutil.rmtree(db_directory)
            st.sidebar.write(f"Directory {db_directory} has been deleted.")
        else:
            st.sidebar.write(f"Directory {db_directory} does not exist or has already been deleted.")
    except Exception as e:
        # Handle any exceptions that occur
        st.sidebar.write(f"An error occurred: {e}")
# -------------------------------------------------------------------------------------
def save_uploaded_file(uploaded_file):
    temp_pdf_path = os.path.join(PDF_FILES_TEMP_FOLDER, f"temp_{uuid.uuid4().hex}.pdf")
    with open(temp_pdf_path, 'wb') as f:
        f.write(uploaded_file.getbuffer())  # Ensure the file is written to the disk
        st.write(temp_pdf_path)
    return temp_pdf_path
# --------------------------------------------------------------------------------------
from sentence_transformers import SentenceTransformer

def sentence_transformer_embedding_function(model_name, query):
    model = SentenceTransformer(model_name)
    return model.encode([query])
# ---------------------------------------------------------------------------------------
def check_duplicate_entry(GET_FILE_NAME):
    check_doc_archive_folder = os.path.join(ARCHIVE_FOLDER, GET_FILE_NAME)
    # Check if the folder exists
    if os.path.exists(check_doc_archive_folder):
        is_duplicate = True
    else:
        is_duplicate = False

    # Return the result of the duplicate check
    return is_duplicate
# ----------------------------------------------------------

def Check_Adding_Chroma(client, collection, combined_data, file_name, file_type):
        
        if file_type == "application/pdf":
            metadata = {}  # For a generic query without specific metadata constraints
            success  = ChromaDB_Process_Add_Data_pdf(client, collection, combined_data)
        else:
            success  = ChromaDB_Process_Add_Data_docx_img(client, collection, combined_data)
        if success:
            st.success(f"Pre proecssing is successful and added to ChromaDB with file name: {file_name}")
            # Archive processed files after successful ChromaDB execution
            archive_processed_files(file_name)
        else:
            st.error("Failed to add data to ChromaDB.")
# -----------------------------------------------------------------------------
import os
import pdfplumber
import uuid
import streamlit as st
SEQUENCE_NUMBER = 1

# ------------------------------------------------------------------------------------
def process_uploaded_file(uploaded_file, client, collection, selected_strategy):
    
    """
    Processes an uploaded file (PDF, DOCX or Images (jpg/png)) and adds the extracted data to a ChromaDB collection.
    It handles file saving, data extraction, and ChromaDB processing. 

    Args:
        uploaded_file (UploadedFile): The file uploaded by the user, provided by Streamlit's file uploader.
        client (ChromaDB.Client): An instance of the ChromaDB client used to interact with the database.
        collection (ChromaDB.Collection): The collection within ChromaDB where the data is to be stored.
        selected_strategy (dict): The strategy used for data extraction, particularly for PDF files.

    Returns:
        None: The function doesn't return anything but performs operations on the uploaded file and ChromaDB.

    Raises:
        Exception: Catches and logs any exceptions that occur during processing.

    Example:
        - For PDF files, it uses `working_extract_data_from_pdf` function to process and extract data.
        - For DOCX files, it uses `extract_data_from_docx`.
        - For IMAGE files, it uses `extract_data_from_img`.
        - After extraction, data is added to ChromaDB using `ChromaDB_Process_Add_Data`.

    Notes:
        - The function saves the uploaded file temporarily for processing and deletes it afterward.
        - It handles different file types and extracts data accordingly.
        - After successful data addition to ChromaDB, it archives processed files.
        - Errors during processing are caught and displayed using Streamlit's error function.
    """
    file_type = uploaded_file.type
    file_name = uploaded_file.name
    temp_file_path = save_uploaded_file(uploaded_file)  # Save the uploaded file temporarily

    if temp_file_path is None:
        st.error("Failed to save the uploaded file.")
        return

    doc_id = f"{file_name}"
    OUTPUT_PATH = os.path.join(PDF_OUTPUT_RESULT_FOLDER, doc_id + ".json")

    try:
        st.write(f'File type: {file_type}')
        if file_type == "application/pdf":
            combined_data = working_extract_data_from_pdf(temp_file_path, OUTPUT_PATH, file_name, selected_strategy)
            if combined_data:
                Check_Adding_Chroma(client, collection, combined_data, file_name, file_type)

        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            combined_data = extract_data_from_docx(temp_file_path, OUTPUT_PATH, file_name)
            #st.write(combined_data)
            if combined_data:
                Check_Adding_Chroma(client, collection, combined_data, file_name, file_type)
        elif uploaded_file.type in ["image/png", "image/jpeg"]:
            combined_data = extract_data_from_img(temp_file_path, OUTPUT_PATH, file_name)
            #st.write(combined_data)
            if combined_data:
                Check_Adding_Chroma(client, collection, combined_data, file_name, file_type)
        else:
            st.error("Unsupported file format")
            return

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
    finally:
        # Remove the temporary file
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            st.write(f"Temporary file {temp_file_path} removed.")

# ---------------------------------------------------------------
def extract_context_from_results(raw_results):
    context = ""
    # Extracting the 'documents' field from raw_results
    documents = raw_results.get('documents', [])
        
    for doc_group in documents:
        for doc in doc_group:
            context += doc + "\n"

    return context
# --------------------------------------------------------------------------------------

from transformers import pipeline

def summarize_text(text):
    model_path_for_summarization = os.path.join(MODEL_FOLDER, 'sshleifer/distilbart-cnn-12-6')
    summarizer = pipeline("summarization", model=model_path_for_summarization)  # Summarization model
    # Consider splitting or truncating the text if it's too long
    summary = summarizer(text[:1024], max_length=130, min_length=30, do_sample=False)  # Flexibile limits
    return summary[0]['summary_text']

# --------------------------------------------------------------------------------------
class GenAIInterface:
    def __init__(self):
        self.LOCAL_ONLY = True
        self.GET_FILE_NAME = None
        self.client = None
        self.collection = None
        self.selected_strategy = None
        self.user_query = "" 
        self.initialize_interface()

        EXTRACT_TABLE_STRATEGY = {
            'vertical_strategy': 'lines',
            'horizontal_strategy': 'text',
            'snap_tolerance': 4
        }
   #-----------------------------------------------------------------------------------
    def initialize_vector_db(self, key_suffix):
        st.title("GenAI Interface")
        selected_vectordb = st.selectbox("Select VectorDB:", options=list(VECTORDB_MAP.keys()), key=f"vector_db_select_{key_suffix}")
        self.client = initialize_client(selected_vectordb)
        heartbeat = self.client.heartbeat()
        st.write(f"Heartbeat: {heartbeat}")
        self.collection = ChromaDB_Initialize(self.client)
        if self.collection:
            st.write("Collection initialized successfully.")
        else:
            st.error("Failed to initialize collection.")
    #-----------------------------------------------------------------------------------
    def fetch_filtered_ids(self, filter_type, filter_value):
  
        """
        Fetches and filters unique IDs (amendment, contract, customer) and dates (validity, load) based on a specified filter from a ChromaDB collection.

        Args:
            filter_type (str): The type of filter to apply (e.g., 'amendment_id', 'contract_id', etc.).
            filter_value (str): The value of the filter to match against the collection.

        Returns:
            dict: A dictionary containing sets of unique IDs and dates filtered by the specified criteria.
                The dictionary keys are 'amendment_ids', 'contract_ids', 'customer_ids', 'validity_dates', and 'load_dates'.

        Raises:
            ValueError: If the collection object is not initialized.
            Exception: Catches and reports any other exceptions during the fetching process.

        Example:
            - If filter_type is 'customer_id' and filter_value is 'C123', the function returns all unique IDs and dates related to 'C123'.

        Notes:
            - The function performs a query on the ChromaDB collection using the provided filter.
            - It extracts and categorizes different types of metadata from the query results.
            - In case of an error during fetching, the function returns empty sets for all categories and displays an error message.
        """
        try:
                # Make sure this line is needed or remove it if not
                #self.initialize_vector_db(key_suffix="main_fetch_unique_ids_def")

                if self.collection is None:
                    raise ValueError("Collection object is not initialized.")

                query_filter = {filter_type: {"$eq": filter_value}}
                query_result = self.collection.query(query_texts=["*"], where=query_filter, include=["metadatas"])

                filtered_amendment_ids, filtered_contract_ids, filtered_customer_ids = set(), set(), set()
                filtered_validity_dates, filtered_load_dates = set(), set()

                for metadata_group in query_result.get("metadatas", []):
                    for metadata in metadata_group:
                        if "amendment_id" in metadata:
                            filtered_amendment_ids.add(metadata["amendment_id"])
                        if "contract_id" in metadata:
                            filtered_contract_ids.add(metadata["contract_id"])
                        if "customer_id" in metadata:
                            filtered_customer_ids.add(metadata["customer_id"])
                        if "validity_date" in metadata:
                            filtered_validity_dates.add(metadata["validity_date"])
                        if "load_date" in metadata:
                            filtered_load_dates.add(metadata["load_date"])
                #st.write(filtered_load_dates)
                return {
                    'amendment_ids': filtered_amendment_ids, 
                    'contract_ids': filtered_contract_ids, 
                    'customer_ids': filtered_customer_ids, 
                    'validity_dates': filtered_validity_dates, 
                    'load_dates': filtered_load_dates
                }
        except Exception as e:
            st.error(f"Error fetching IDs: {str(e)}")
            return set(), set(), set(), set(), set()

# ------------------------------------------------------------------------------
    from datetime import datetime

    def fetch_unique_ids(self):
        try:
            if self.collection is None:
                raise ValueError("Collection object is not initialized.")

            # Query the collection
            query_result = self.collection.query(query_texts=["*"], include=["metadatas"])

            # Initialize sets to store unique IDs
            unique_contract_ids = set()
            unique_customer_ids = set()
            unique_amendment_ids = set()
            unique_load_dates = set()
            unique_source_ids = set()
            unique_validity_dates = set()

            # Process each metadata entry
            for metadata_group in query_result.get("metadatas", []):
                for metadata in metadata_group:
                    contract_id = metadata.get("contract_id")
                    customer_id = metadata.get("customer_id")
                    amendment_id = metadata.get("amendment_id")
                    load_date = metadata.get("load_date", "")
                    source_id = metadata.get("source")
                    validity_date = metadata.get("validity_date")

                    # Add values to respective sets
                    if contract_id:
                        unique_contract_ids.add(contract_id)
                    if customer_id:
                        unique_customer_ids.add(customer_id)
                    if amendment_id:
                        unique_amendment_ids.add(amendment_id)
                    if load_date:
                        unique_load_dates.add(load_date)
                    if source_id:
                        unique_source_ids.add(source_id)
                    if validity_date:
                        unique_validity_dates.add(validity_date)


            # Return a dictionary containing all the unique sets
            return {
                'contract_ids': unique_contract_ids,
                'customer_ids': unique_customer_ids,
                'amendment_ids': unique_amendment_ids,
                'load_dates': unique_load_dates,
                'source_ids': unique_source_ids,
                'validity_dates': unique_validity_dates
            }

        except Exception as e:
            st.error(f"Error fetching IDs: {str(e)}")
            return None
# ---------------------------------------------------------------------------
    def get_document_text_for_source_id(self, collection, source_id):
        results = collection.query(query_texts=["*"], where={"source": {"$eq": source_id}}, include=["documents"])
        document_text = ""
        for group in results.get("documents", []):
            for document in group:
                document_text += document + "\n"
        return document_text
# --------------------------------------------------------------------------------
    def execute_query(self, user_query, selected_contract_id, selected_customer_id, selected_amendment_id, selected_load_date, selected_source_id, additional_filter, NUMBER_SEARCH_QUERY, selected_validity_date, use_llm=None, selected_model=None, llm_choice=None, enable_summarization=False, summarize_source_id='None'):
        """
        Executes a query on a ChromaDB collection and optionally processes the query with a Language Learning Model (LLM).

        Args:
            user_query (str): The query text to be executed.
            selected_contract_id (str): The selected contract ID for filtering.
            selected_customer_id (str): The selected customer ID for filtering.
            selected_amendment_id (str): The selected amendment ID for filtering.
            selected_load_date (str): The selected load date for filtering.
            selected_source_id (str): The selected source ID for filtering.
            additional_filter (str): Additional filters to be applied to the query.
            NUMBER_SEARCH_QUERY (int): The number of results to return.
            selected_validity_date (str): The selected validity date for filtering.
            use_llm (bool): Flag to determine if the LLM should be used for processing the query.
            selected_model (str): The selected LLM model.
            llm_choice (str): The choice of LLM (e.g., OpenLLM).
            enable_summarization (bool): Flag to enable summarization of results.
            summarize_source_id (str): The source ID to be used for summarization.

        Returns:
            None: The function directly interacts with Streamlit UI and does not return any value.

        Raises:
            ValueError: If the collection object is not initialized.

        Example:
            execute_query("find contract details", "C123", None, None, None, None, "extra filter", 5, None)

        Notes:
            - The function queries the ChromaDB collection based on the provided parameters.
            - It supports complex filtering and can combine multiple conditions.
            - If `use_llm` is True and a model is loaded, the function processes the query with the LLM.
            - In case of any error during query execution or LLM processing, an appropriate error message is displayed.
        """

        # Ensure the 'collection' object is initialized
        if self.collection is None:
            st.error("Collection object is not initialized.")
            raise ValueError("Collection object is not initialized.")
        
        # Parse additional filters
        additional_filter_dict = self.parse_additional_filter(additional_filter)
        
        # Initialize query parameters
        query_params = {
            'query_texts': [user_query],
            'n_results': NUMBER_SEARCH_QUERY,
            'where': {}
        }

        # Build where clause based on selected filters
        where_conditions = []
        if selected_contract_id and selected_contract_id != 'None':
            where_conditions.append({'contract_id': {"$in": [selected_contract_id]}})
        if selected_customer_id and selected_customer_id != 'None':
            where_conditions.append({'customer_id': {"$in": [selected_customer_id]}})
        if selected_amendment_id and selected_amendment_id != 'None':
            where_conditions.append({'amendment_id': {"$in": [selected_amendment_id]}})
        if selected_validity_date and selected_validity_date != 'None':
            where_conditions.append({'validity_date': {"$eq": selected_validity_date}})
        if selected_source_id and selected_source_id != 'None':
            where_conditions.append({'source': {"$eq": selected_source_id}})

        # Incorporate additional where_document filter if provided
        if additional_filter_dict:
            query_params['where_document'] = additional_filter_dict

        # Determine logical operator based on the number of conditions
        if len(where_conditions) > 1:
            query_params['where'] = {'$or': where_conditions}
        elif where_conditions:
            query_params['where'] = where_conditions[0]

        # Perform the query
        st.write(query_params)
        try:
            raw_results = self.collection.query(**query_params)
            st.write(raw_results)
        except Exception as e:
            st.error(f"Query execution failed: {str(e)}")
            return

        if raw_results:
            try:
                if use_llm and selected_model != 'None' and 'loaded_model' in st.session_state:
                    context = extract_context_from_results(raw_results)
                    #st.write(f'Printing before llm response {context}')
                    llm_response = process_query_with_llm(self.user_query, context, self.selected_model, self.llm_choices == 'OpenLLM')
                    st.write(f'llm response : {llm_response}')
                    st.write(llm_response)
                elif use_llm and selected_model != 'None':
                    st.warning("Model is still loading or not loaded. Please wait or reload the model.")
            except Exception as e:
                st.error(f"Error processing LLM response: {str(e)}")
        else:
            st.write("No results found for the query.")

# ----------------------------------------------------------------------------------   
    def handle_summarization(self, summarize_source_id):
        try:
            if not summarize_source_id or summarize_source_id == 'None':
                st.error("Please select a valid Source ID for summarization.")
                return

            document_text = self.get_document_text_for_source_id(self.collection, summarize_source_id)
            if not document_text.strip():
                st.error("No text available for summarization.")
            else:
                # Split the text into chunks of 1024 characters
                chunks = [document_text[i:i+1024] for i in range(0, len(document_text), 1024)]

                # Summarize each chunk and concatenate summaries
                summarized_text = ""
                for chunk in chunks:
                    summary_chunk = summarize_text(chunk)
                    summarized_text += summary_chunk + "\n\n"

                st.write("Summarized Text:")
                st.text_area("Summary", summarized_text, height=250)
                st.download_button(label="Download Summary", data=summarized_text, file_name="summarized_text.txt", mime="text/plain")

        except Exception as e:
            st.error(f"Summarization error: {str(e)}")

    # ----------------------------------------------------------------------------
    def initialize_interface(self):
        try:
            self.initialize_vector_db(key_suffix="main_initialize_interface")
            if self.collection is None:
                raise ValueError("Collection object is not initialized.")

            # Model Selection and LLM Options
            model_choices_with_none = ['None'] + list(model_choices.keys())
            self.selected_model = st.selectbox("Select a model:", model_choices_with_none)
            self.use_llm = st.checkbox("Use Large Language Model for Enhanced Results")
            self.llm_choices = st.radio("Choose the LLM approach", ('OpenLLM', 'Transformers'))
            st.markdown(f'**Your LLM Choice:** {self.llm_choices}')
            self.user_query = st.text_input("Enter your query:")
            st.markdown(f'**Your Question:** {self.user_query}')

            # Fetch unique IDs for filtering
            unique_ids = self.fetch_unique_ids()
 
            # Add summarization options
            self.enable_summarization = st.checkbox("Enable Summarization")

            if self.enable_summarization:
                # Allow users to select a source ID for summarization
                self.summarize_source_id = st.selectbox("Select Source ID for Summarization:", ['None'] + list(unique_ids['source_ids']))

                # Summarization Button
                if st.button("Generate Summary"):
                    self.handle_summarization(summarize_source_id=self.summarize_source_id)

            selected_source_id = st.selectbox("Select Source ID:", ['None'] + list(unique_ids['source_ids']))
            selected_contract_id = st.selectbox("Select Contract ID:", ['None'] + list(unique_ids['contract_ids']))
            selected_customer_id = st.selectbox("Select Customer ID:", ['None'] + list(unique_ids['customer_ids']))

            # More filter options
            selected_amendment_id = st.selectbox("Select Amendment ID:", ['None'] + list(unique_ids.get('amendment_ids', [])))
            selected_validity_date = st.selectbox("Select Validity Date:", ['None'] + list(unique_ids.get('validity_dates', [])))
            selected_load_date = st.selectbox("Select Load Date:", ['None'] + list(unique_ids.get('load_dates', [])))
            additional_filter = st.text_input("Enter additional filter criteria:")
            NUMBER_SEARCH_QUERY = st.sidebar.number_input("Number of Results to Fetch", min_value=1, value=1)

            # Logical operator buttons
            logic_ops1, logic_ops2 = st.columns(2)
            with logic_ops1:
                and_selected = st.button("AND")
            with logic_ops2:
                or_selected = st.button("OR")
            logical_operator = "AND" if and_selected else "OR"

            # Load Model Button
            if st.button("Load Model"):
                # Trigger asynchronous model loading
                load_llm_async(self.selected_model, self.llm_choices == 'OpenLLM', load_llm_callback)
            # Execute Query Button
            if st.button("Execute Query"):
                self.execute_query(
                    user_query=self.user_query,
                    selected_contract_id=selected_contract_id,
                    selected_customer_id=selected_customer_id,
                    selected_amendment_id=selected_amendment_id,
                    selected_load_date=selected_load_date,
                    selected_source_id=selected_source_id,
                    additional_filter=additional_filter,
                    NUMBER_SEARCH_QUERY=NUMBER_SEARCH_QUERY,
                    selected_validity_date=selected_validity_date,
                    use_llm=self.use_llm,
                    selected_model=self.selected_model,
                    llm_choice=self.llm_choices
                )

        except Exception as e:
            st.error(f"Error in initialize_interface: {str(e)}")
#-----------------------------------------------------------------------------------
    def handle_reset_database(self):
        if st.button("Confirm Reset"):
            if self.client is not None:
                reset_database(self.client, VECTORDB_DIR_SPACE_FOLDER)
                st.success("Database reset successfully.")
            else:
                st.error("Client not initialized. Cannot reset database.")
#-----------------------------------------------------------------------------------
    def handle_archive_data(self):
        if st.button("Confirm Archive"):
            files_in_folder = os.listdir(PDF_OUTPUT_RESULT_FOLDER)
            if not files_in_folder:
                st.write("Nothing to archive.")
            else:
                for selected_file in files_in_folder:
                    self.GET_FILE_NAME = os.path.splitext(selected_file)[0]
                    archive_processed_files(self.GET_FILE_NAME)
                    st.success(f"Archived file: {self.GET_FILE_NAME}")

                # Optionally, list archived files here
                st.write("Archived files:")
                for file in files_in_folder:
                    st.write(file)
#--------------------------------------------------------------------------------------------
    def execute_vector_db(self):
        #selected_vectordb = st.selectbox("Select VectorDB:", options=["ChromaDB", "Faiss", "Scann"])
        if self.selected_vectordb == "ChromaDB":
            success = ChromaDB_Process_Add_Data_pdf(self.client, self.collection, self.GET_FILE_NAME)
        elif self.selected_vectordb == "Faiss":
            pass  # Add Faiss-specific initialization if needed
        elif self.selected_vectordb == "Scann":
            pass  # Add Scann-specific initialization if needed
#-----------------------------------------------------------------------------------    
    def select_strategy(self):
        selected_task = st.selectbox("Choose Task", list(ENGINE_LISTS.values()))
        default_strategy_index = list(EXTRACTION_STRATEGIES.keys()).index('text_lines') if 'text_lines' in EXTRACTION_STRATEGIES else 0
        self.selected_strategy = st.sidebar.selectbox("Choose a strategy", options=list(EXTRACTION_STRATEGIES.keys()), index=default_strategy_index)

    def process_uploaded_files(self, uploaded_files):
        if uploaded_files:
            for uploaded_file in uploaded_files:
                process_uploaded_file(uploaded_file, self.client, self.collection, self.selected_strategy)

    # ------------------------------------------------------------------
    def parse_additional_filter(self, filter_string):
        if filter_string:
            return {"$contains": filter_string}
        return {}

    # --------------- main function -------------------------------------------------
    def main(self):
        with st.sidebar:
            st.title("Settings")
        
            # Existing code for actions
            action = st.selectbox("Choose Action:", ["None", "Reset Database", "Archive Data"])
            if action == "Reset Database":
                self.handle_reset_database()
            elif action == 'Archive Data':
                self.handle_archive_data()
            
            self.select_strategy()

            # Updated file uploader to accept PDF, DOCX, PNG, and JPG files
            uploaded_files = st.file_uploader("Upload your files", type=["pdf", "docx", "png", "jpg"], accept_multiple_files=True, key="file_uploader_sidebar")
            self.process_uploaded_files(uploaded_files)
#-----------------------------------------------------------------------------------
if __name__ == "__main__":
    
    gen_ai_interface = GenAIInterface()
    gen_ai_interface.main()